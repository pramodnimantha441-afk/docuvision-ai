from fastapi import APIRouter, Depends, HTTPException, Response
from fastapi.responses import StreamingResponse
from firebase_admin import db
from middleware.auth_middleware import verify_firebase_token
from models.schemas import DocumentSaveRequest
from datetime import datetime
import uuid
import io
import re
import base64
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

router = APIRouter()

@router.get('/documents')
async def get_documents(token_data: dict = Depends(verify_firebase_token)):
    uid = token_data['uid']
    ref = db.reference(f'/documents/{uid}')
    docs = ref.get() or {}
    result = [{'id': k, **v} for k, v in docs.items()]
    result.sort(key=lambda x: x.get('createdAt', 0), reverse=True)
    return result

@router.post('/documents/save')
async def save_document(
    doc: DocumentSaveRequest,
    token_data: dict = Depends(verify_firebase_token)
):
    uid = token_data['uid']
    doc_id = str(uuid.uuid4())
    now = datetime.utcnow().timestamp() * 1000
    ref = db.reference(f'/documents/{uid}/{doc_id}')
    data = doc.dict()
    data['createdAt'] = now
    data['updatedAt'] = now
    data['userId'] = uid
    ref.set(data)
    return {'id': doc_id, 'message': 'Saved successfully', **data}

@router.delete('/documents/{doc_id}')
async def delete_document(
    doc_id: str,
    token_data: dict = Depends(verify_firebase_token)
):
    uid = token_data['uid']
    ref = db.reference(f'/documents/{uid}/{doc_id}')
    if ref.get() is None:
        raise HTTPException(status_code=404, detail='Document not found')
    ref.delete()
    return {'message': 'Deleted successfully'}

@router.put('/documents/{doc_id}')
async def update_document(
    doc_id: str,
    data: dict,
    token_data: dict = Depends(verify_firebase_token)
):
    uid = token_data['uid']
    ref = db.reference(f'/documents/{uid}/{doc_id}')
    if ref.get() is None:
        raise HTTPException(status_code=404, detail='Document not found')
    from datetime import datetime
    data['updatedAt'] = datetime.utcnow().timestamp() * 1000
    ref.update(data)
    return {'message': 'Updated successfully'}


def strip_html_tags(html_str: str) -> str:
    """Convert HTML from Quill editor into clean readable paragraphs."""
    if not html_str:
        return ""
    text = re.sub(r'</?(p|br|div|li|h[1-6])[^>]*>', '\n', html_str)
    text = re.sub(r'<[^>]+>', '', text)
    text = text.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    return '\n'.join(lines)


@router.post('/documents/export-docx')
async def export_docx(
    payload: dict
):
    """
    Generate a 100% compliant, beautifully formatted Microsoft Word (.docx) file.
    """
    try:
        title = payload.get('title') or 'DocuVision AI Digitized Document'
        raw_text = payload.get('text') or ''
        summary = payload.get('summary') or ''
        key_points = payload.get('key_points') or []
        action_items = payload.get('action_items') or []
        image_base64 = payload.get('image_base64') or ''

        clean_text = strip_html_tags(raw_text)

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # 1. Header
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = header_p.add_run("DocuVision AI • Document Digitization Workspace")
        hr.font.name = "Calibri"
        hr.font.size = Pt(8.5)
        hr.font.color.rgb = RGBColor(120, 144, 156)

        # 2. Main Title
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(6)
        title_p.paragraph_format.space_after = Pt(4)
        tr = title_p.add_run(str(title))
        tr.font.name = "Calibri"
        tr.font.size = Pt(20)
        tr.bold = True
        tr.font.color.rgb = RGBColor(26, 35, 126)

        # Metadata subtitle
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(14)
        date_str = datetime.now().strftime("%B %d, %Y - %I:%M %p")
        mr = meta_p.add_run(f"Digitized on {date_str}  |  Status: Verified Complete")
        mr.font.name = "Calibri"
        mr.font.size = Pt(9.5)
        mr.italic = True
        mr.font.color.rgb = RGBColor(90, 100, 110)

        # 3. AI Executive Summary
        if summary and str(summary).strip():
            sum_h = doc.add_paragraph()
            sum_h.paragraph_format.space_before = Pt(8)
            sum_h.paragraph_format.space_after = Pt(4)
            shr = sum_h.add_run("1.0 AI Executive Summary")
            shr.font.name = "Calibri"
            shr.font.size = Pt(13)
            shr.bold = True
            shr.font.color.rgb = RGBColor(40, 53, 147)

            try:
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                cell = table.cell(0, 0)
                cell.width = Inches(6.5)

                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4FF"/>')
                cell._tc.get_or_add_tcPr().append(shd)

                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    '<w:left w:val="single" w:sz="24" w:space="0" w:color="3F51B5"/>'
                    '<w:top w:val="none"/>'
                    '<w:right w:val="none"/>'
                    '<w:bottom w:val="none"/>'
                    '</w:tcBorders>'
                )
                tcPr.append(tcBorders)

                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(6)
                cp.paragraph_format.space_after = Pt(6)
                cp.paragraph_format.left_indent = Inches(0.15)
                cp.paragraph_format.right_indent = Inches(0.15)
                cpr = cp.add_run(str(summary).strip())
                cpr.font.name = "Calibri"
                cpr.font.size = Pt(10.5)
                cpr.font.color.rgb = RGBColor(33, 33, 33)
                doc.add_paragraph()
            except Exception as e:
                sp = doc.add_paragraph(str(summary).strip())
                sp.paragraph_format.space_after = Pt(8)

        # 4. Key Points
        if isinstance(key_points, list) and len(key_points) > 0:
            kp_h = doc.add_paragraph()
            kp_h.paragraph_format.space_before = Pt(6)
            kp_h.paragraph_format.space_after = Pt(4)
            kpr = kp_h.add_run("Key Highlights")
            kpr.font.name = "Calibri"
            kpr.font.size = Pt(11)
            kpr.bold = True

            for point in key_points:
                pt_p = doc.add_paragraph(style='List Bullet')
                pt_p.paragraph_format.space_after = Pt(2)
                ptr = pt_p.add_run(str(point))
                ptr.font.name = "Calibri"
                ptr.font.size = Pt(10)

            doc.add_paragraph()

        # 5. Transcribed Content
        tx_h = doc.add_paragraph()
        tx_h.paragraph_format.space_before = Pt(10)
        tx_h.paragraph_format.space_after = Pt(6)
        txr = tx_h.add_run("2.0 Extracted Handwritten Transcription")
        txr.font.name = "Calibri"
        txr.font.size = Pt(13)
        txr.bold = True
        txr.font.color.rgb = RGBColor(40, 53, 147)

        for line in clean_text.split('\n'):
            if line.strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                pr = p.add_run(line.strip())
                pr.font.name = "Calibri"
                pr.font.size = Pt(11)
                pr.font.color.rgb = RGBColor(33, 33, 33)

        # 6. Optional: Embed original image
        if image_base64 and len(image_base64) > 100:
            try:
                if ',' in image_base64:
                    b64_data = image_base64.split(',', 1)[1]
                else:
                    b64_data = image_base64
                img_bytes = base64.b64decode(b64_data)
                img_stream = io.BytesIO(img_bytes)

                doc.add_paragraph()
                img_h = doc.add_paragraph()
                img_h.paragraph_format.space_before = Pt(14)
                img_h.paragraph_format.space_after = Pt(6)
                ihr = img_h.add_run("Appendix: Original Document Source")
                ihr.font.name = "Calibri"
                ihr.font.size = Pt(12)
                ihr.bold = True
                ihr.font.color.rgb = RGBColor(90, 100, 110)

                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.add_run().add_picture(img_stream, width=Inches(5.5))
            except Exception as e:
                print(f"Warning: could not embed image in docx: {e}")

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = f"{re.sub(r'[^a-zA-Z0-9_-]', '_', str(title))}_{datetime.now().strftime('%Y%m%d')}.docx"

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as exc:
        print(f"Error generating DOCX: {exc}")
        raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(exc)}")
